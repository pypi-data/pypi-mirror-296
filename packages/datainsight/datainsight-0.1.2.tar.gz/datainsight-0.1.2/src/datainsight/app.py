import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import LabelEncoder
import numpy as np
import os, io
import shutil, random
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ftplib

# Define FTP server details
FTP_SERVER = 'ftp.chears.in'
FTP_USER = 'webpush@chears.in'
FTP_PASS = 'Web@12358'
UPLOAD_DIRECTORY = '/uploads'  # The directory on the FTP server


def save_uploaded_file(uploaded_file, email):
    """ Save the uploaded file to the FTP server """
    n = random.randint(0, 100)
    with ftplib.FTP(FTP_SERVER) as ftp:
        ftp.login(FTP_USER, FTP_PASS)
        ftp.cwd(UPLOAD_DIRECTORY)
        with io.BytesIO(uploaded_file.getvalue()) as file:
            ftp.storbinary(
                f'STOR useremail_{email.replace("@", "_at_").replace(".", "_")}_{n}_{uploaded_file.name}',
                file
            )

# Connect to the FTP server
ftp = ftplib.FTP(FTP_SERVER)
ftp.login(FTP_USER, FTP_PASS)

# Retrieve the file
def retrieve_file(file_path):
    with io.BytesIO() as file_buffer:
        ftp.retrbinary(f"RETR {file_path}", file_buffer.write)
        file_buffer.seek(0)  # Go to the beginning of the file
        return file_buffer.read()

# Read the file into a pandas DataFrame
user_data = retrieve_file("/data/users.csv")
user_data = pd.read_csv(io.StringIO(user_data.decode('utf-8')))
print(user_data)

TEMP_CHARTS_DIR = "temp_charts"
WATERMARK_TEXT = "Developed by datainsight (CHEARS)"


def load_file(file_like):
    """Load CSV or Excel file from a file-like object"""
    if file_like.name.endswith('.csv'):
        df = pd.read_csv(file_like)
    elif file_like.name.endswith('.xls') or file_like.name.endswith('.xlsx'):
        sheets = pd.ExcelFile(file_like).sheet_names
        sheet_name = st.selectbox("Select sheet", sheets)
        df = pd.read_excel(file_like, sheet_name=sheet_name)
    else:
        st.error("Unsupported file type.")
        return None
    return df

def impute_missing_values(df):
    impute_method = st.selectbox(
        "Select Imputation Method",
        options=['None', 'Mean', 'Median', 'Most Frequent', 'Constant']
    )
    if impute_method != 'None':
        if impute_method == 'Constant':
            constant_value = st.number_input("Enter Constant Value", value=0)
            imputer = SimpleImputer(strategy='constant', fill_value=constant_value)
        else:
            imputer = SimpleImputer(strategy=impute_method.lower())

        for column in df.columns:
            if df[column].isnull().any():
                if pd.api.types.is_numeric_dtype(df[column]):
                    df[column] = imputer.fit_transform(df[[column]])
                elif df[column].dtype == 'object':
                    df[column] = imputer.fit_transform(df[[column]].astype(str))

        st.write("### Updated Dataset Overview")
        st.write(df.info())
        st.write(df.describe(include='all'))

    return df

def advanced_processing(df):
    st.write("### Advanced Processing")

    # Convert nominal or string columns to numerical
    string_columns = df.select_dtypes(include=['object']).columns
    for column in string_columns:
        st.write(f"#### Processing Column: {column}")
        method = st.selectbox(
            f"Select Encoding Method for {column}",
            options=['Label Encoding', 'One Hot Encoding', 'Ordinal Encoding']
        )
        if method == 'Label Encoding':
            le = LabelEncoder()
            df[column] = le.fit_transform(df[column].astype(str))
            st.write(f"Column {column} has been label encoded.")
        elif method == 'One Hot Encoding':
            df = pd.get_dummies(df, columns=[column])
            st.write(f"Column {column} has been one hot encoded.")
        elif method == 'Ordinal Encoding':
            unique_values = df[column].unique()
            sorted_values = sorted(unique_values)
            mapping = {val: i for i, val in enumerate(sorted_values)}
            df[column] = df[column].map(mapping)
            st.write(f"Column {column} has been ordinally encoded.")

    # Convert numerical columns to bins or categories
    if st.checkbox("Convert Numerical Columns to Bins/Categories"):
        for column in df.select_dtypes(include=[np.number]).columns:
            st.write(f"#### Binning Column: {column}")
            bins = st.slider(f"Select number of bins for {column}", min_value=2, max_value=20, value=5)
            bin_edges = np.linspace(df[column].min(), df[column].max(), bins + 1)
            df[f'{column}_binned'] = pd.cut(df[column], bins=bin_edges, labels=[f'Bin_{i}' for i in range(1, bins + 1)])
            st.write(f"Column {column} has been binned into {bins} categories.")

    st.write("### Updated Dataset Overview")
    st.write(df.info())
    st.write(df.describe(include='all'))

    return df

def correct_data_types(df):
    st.write("### Correct Data Types")

    # Create a table for data type conversion with dropdowns for each column
    conversion_results = []
    columns = df.columns.tolist()
    current_types = df.dtypes.astype(str).tolist()
    
    conversion_options = ['None', 'int', 'float', 'str', 'datetime']
    new_types = []

    st.write("#### Data Type Conversion Table")

    # Show current data types, dropdown to change, and result/remark in a single table
    for idx, column in enumerate(columns):
        st.write(f"Processing Column: {column}")
        col1, col2, col3, col4 = st.columns([2, 2, 2, 4])

        with col1:
            col1.write(column)
        with col2:
            col2.write(current_types[idx])
        with col3:
            selected_type = col3.selectbox(f"Select new type for {column}", options=conversion_options, key=column)
            new_types.append(selected_type)
        with col4:
            remark = "No change"
            if selected_type != 'None' and selected_type != current_types[idx]:
                try:
                    if selected_type == 'int':
                        df[column] = df[column].astype(int)
                    elif selected_type == 'float':
                        df[column] = df[column].astype(float)
                    elif selected_type == 'str':
                        df[column] = df[column].astype(str)
                    elif selected_type == 'datetime':
                        df[column] = pd.to_datetime(df[column])
                    remark = f"Successfully converted to {selected_type}"
                except Exception as e:
                    remark = f"Failed: {e}"

            col4.write(remark)

    # Show the updated data types and result table
    st.write("### Updated Dataset Overview")
    st.write(df.info())

    return df


def create_chart(df, col_x, col_y, chart_type, file_path):
    """ Create a chart and save it to file_path """
    fig, ax = plt.subplots(figsize=(10, 6))
    
    if chart_type == 'Scatter Plot':
        ax.scatter(df[col_x], df[col_y])
        ax.set_xlabel(col_x)
        ax.set_ylabel(col_y)
        ax.set_title(f"Scatter Plot of {col_x} vs {col_y}")

    elif chart_type == 'Line Plot':
        ax.plot(df[col_x], df[col_y])
        ax.set_xlabel(col_x)
        ax.set_ylabel(col_y)
        ax.set_title(f"Line Plot of {col_x} vs {col_y}")

    elif chart_type == 'Bar Plot':
        df.groupby(col_x)[col_y].mean().plot(kind='bar', ax=ax)
        ax.set_xlabel(col_x)
        ax.set_ylabel(col_y)
        ax.set_title(f"Bar Plot of {col_x} by {col_y}")

    elif chart_type == 'Histogram':
        df[col_x].plot(kind='hist', ax=ax, bins=30)
        ax.set_xlabel(col_x)
        ax.set_title(f"Histogram of {col_x}")

    elif chart_type == 'Box Plot':
        sns.boxplot(x=df[col_x], y=df[col_y], ax=ax)
        ax.set_xlabel(col_x)
        ax.set_ylabel(col_y)
        ax.set_title(f"Box Plot of {col_x} vs {col_y}")

    elif chart_type == 'Heatmap':
        sns.heatmap(df.corr(), annot=True, fmt=".2f", cmap='coolwarm', ax=ax)
        ax.set_title('Correlation Heatmap')

    elif chart_type == 'Pie Chart':
        df[col_x].value_counts().plot.pie(autopct='%1.1f%%', ax=ax)
        ax.set_ylabel('')
        ax.set_title(f"Pie Chart of {col_x}")

    elif chart_type == 'Trend View':
        ax.plot(df[col_x], df[col_y], label=f"Trend of {col_y} over {col_x}")
        ax.set_xlabel(col_x)
        ax.set_ylabel(col_y)
        ax.set_title(f"Trend View: {col_y} over {col_x}")

    elif chart_type == 'Seasonality':
        df['Month'] = df[col_x].dt.month
        df.groupby('Month')[col_y].mean().plot(ax=ax)
        ax.set_xlabel('Month')
        ax.set_ylabel(col_y)
        ax.set_title(f"Seasonality: {col_y} by Month")
    
    fig.text(0.95, 0.05, WATERMARK_TEXT, fontsize=8, color='gray', ha='right', va='bottom')
    plt.savefig(file_path)
    plt.close()

    # Add watermark to the saved chart
    # add_watermark_to_chart(file_path, file_path)

def add_table_to_doc(doc, df):
    """Add a DataFrame as a table to the Word document"""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
    
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc

def generate_full_report(df, selected_columns):
    """Generates a full report in a Word document and provides a download link."""
    if not os.path.exists(TEMP_CHARTS_DIR):
        os.makedirs(TEMP_CHARTS_DIR)
    
    doc = Document()
    doc.add_heading('Full Data Report', 0)
    
    # 1. Data Summary Table
    data_summary = pd.DataFrame({
        'Column Name': df.columns,
        'Data Type': df.dtypes.astype(str),
        'Number of Values': df.notna().sum(),
        'Number of Unique Values': df.nunique(),
        'Number of Null Values': df.isna().sum()
    })
    
    doc.add_heading('Data Summary Table', level=1)
    add_table_to_doc(doc, data_summary)

    # 2. Description Table for Numerical Values
    num_desc = df.describe(include=[np.number])
    doc.add_heading('Numerical Values Description', level=1)
    add_table_to_doc(doc, num_desc)

    # 3. Charts
    for col in selected_columns:
        if df[col].dtype == 'object':
            create_chart(df, col, None, 'Pie Chart', os.path.join(TEMP_CHARTS_DIR, f'{col}_pie_chart.png'))
        elif pd.api.types.is_numeric_dtype(df[col]):
            create_chart(df, col, None, 'Histogram', os.path.join(TEMP_CHARTS_DIR, f'{col}_histogram.png'))
            if len(df.columns) > 1:
                for other_col in df.columns:
                    if other_col != col and pd.api.types.is_numeric_dtype(df[other_col]):
                        create_chart(df, col, other_col, 'Scatter Plot', os.path.join(TEMP_CHARTS_DIR, f'{col}_vs_{other_col}_scatter_plot.png'))
                        create_chart(df, col, other_col, 'Line Plot', os.path.join(TEMP_CHARTS_DIR, f'{col}_vs_{other_col}_line_plot.png'))
                        create_chart(df, col, other_col, 'Bar Plot', os.path.join(TEMP_CHARTS_DIR, f'{col}_vs_{other_col}_bar_plot.png'))
    
    for col_x in selected_columns:
        for col_y in selected_columns:
            if col_x != col_y and pd.api.types.is_numeric_dtype(df[col_x]) and pd.api.types.is_numeric_dtype(df[col_y]):
                create_chart(df, col_x, col_y, 'Scatter Plot', os.path.join(TEMP_CHARTS_DIR, f'{col_x}_vs_{col_y}_scatter_plot.png'))
                create_chart(df, col_x, col_y, 'Line Plot', os.path.join(TEMP_CHARTS_DIR, f'{col_x}_vs_{col_y}_line_plot.png'))
                create_chart(df, col_x, col_y, 'Bar Plot', os.path.join(TEMP_CHARTS_DIR, f'{col_x}_vs_{col_y}_bar_plot.png'))

    # Add charts to the document
    for image_file in os.listdir(TEMP_CHARTS_DIR):
        doc.add_heading(image_file, level=2)
        doc.add_picture(os.path.join(TEMP_CHARTS_DIR, image_file), width=Inches(6))

    # Cleanup
    shutil.rmtree(TEMP_CHARTS_DIR)

    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def main():
    st.title("Data Analysis and Processing")
    email = st.text_input("Enter registered email:")
    if email:
        if email in user_data.email.to_list():
            tab1, tab2 = st.tabs(["Data Analysis", "Full Report"])
            uploaded_file = st.file_uploader("Upload your CSV or Excel file", type=["csv", "xlsx", "xls"])
            if uploaded_file is not None:
                file_path = save_uploaded_file(uploaded_file, email)
                df = load_file(uploaded_file)
                if df is not None:
                    with tab1:
                        st.write("### Initial Dataset Overview")
                        st.write(df.info())
                        st.write(df.describe(include='all'))

                        # Step 1: Impute missing values
                        df = impute_missing_values(df)

                        # Step 2: Perform advanced processing
                        df = advanced_processing(df)

                        # Step 3: Correct data types
                        df = correct_data_types(df)

                        st.write("### Data Preprocessing Complete")

                        # Step 4: Ask for chart types and columns
                        chart_type = st.selectbox(
                            "Select Chart Type",
                            options=[
                                'Scatter Plot', 'Line Plot', 'Bar Plot', 'Histogram', 'Box Plot',
                                'Heatmap', 'Pie Chart', 'Trend View', 'Seasonality'
                            ]
                        )

                        numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
                        timestamp_columns = df.select_dtypes(include=[np.datetime64]).columns.tolist()
                        categorical_columns = df.select_dtypes(include=['object']).columns.tolist()

                        # Based on chart type, ask for the required columns
                        if chart_type in ['Scatter Plot', 'Line Plot', 'Bar Plot', 'Box Plot', 'Trend View']:
                            col_x = st.selectbox("Select X-axis Column", options=numeric_columns + timestamp_columns)
                            col_y = st.selectbox("Select Y-axis Column", options=numeric_columns)
                            if col_x and col_y:
                                with st.empty():
                                    create_chart(df, col_x, col_y, chart_type, "chart.png")
                                    st.image("chart.png")

                        elif chart_type == 'Histogram':
                            col_x = st.selectbox("Select Column for Histogram", options=numeric_columns)
                            if col_x:
                                with st.empty():
                                    create_chart(df, col_x, None, chart_type, "chart.png")
                                    st.image("chart.png")

                        elif chart_type == 'Pie Chart':
                            col_x = st.selectbox("Select Categorical Column for Pie Chart", options=categorical_columns)
                            if col_x:
                                with st.empty():
                                    create_chart(df, col_x, None, chart_type, "chart.png")
                                    st.image("chart.png")

                        elif chart_type == 'Heatmap':
                            if len(numeric_columns) > 1:
                                with st.empty():
                                    create_chart(df, None, None, chart_type, "chart.png")
                                    st.image("chart.png")

                        elif chart_type == 'Seasonality':
                            if timestamp_columns:
                                col_x = st.selectbox("Select Time Column", options=timestamp_columns)
                                col_y = st.selectbox("Select Column for Seasonality", options=numeric_columns)
                                if col_x and col_y:
                                    df[col_x] = pd.to_datetime(df[col_x])
                                    with st.empty():
                                        create_chart(df, col_x, col_y, chart_type, "chart.png")
                                        st.image("chart.png")

                    with tab2:
                        selected_columns = st.multiselect("Select Columns for Full Report", options=df.columns)
                        print(selected_columns)
                        if selected_columns:
                            if st.button("Generate Full Report"):
                                report_io = generate_full_report(df, selected_columns)
                                st.download_button(
                                    label="Download Full Report",
                                    data=report_io,
                                    file_name="full_report.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        else:
                            st.error("Please select at least one column for analysis.")
        else:
            st.error("You are not authorized to use this service. Please write email to conatct@chears.in to request for access.")

    else:
        st.write("Please enter email.")
if __name__ == "__main__":
    main()
